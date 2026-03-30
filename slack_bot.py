"""
タビ男 - Slack AI社員
機能:
  - @メンションで質問に回答（Claude AI）
  - 日報を自動検出 → GoogleスプレッドシートにKPI記録
  - 契約書・雛形をGoogle Driveから検索してSlackに直接投稿
  - 毎朝9時に前日の進捗サマリーを投稿
"""

import io
import os
import re
import json
import logging
from slack_bolt import App
from slack_bolt.adapter.socket_mode import SocketModeHandler
from anthropic import Anthropic
from google.oauth2 import service_account
from googleapiclient.discovery import build
from apscheduler.schedulers.background import BackgroundScheduler
from dotenv import load_dotenv

load_dotenv()
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")

# ─── クライアント初期化 ────────────────────────────────────────────────────────

slack_app = App(token=os.environ["SLACK_BOT_TOKEN"])
claude = Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])

GOOGLE_SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive"
]

def _get_google_creds():
    creds_env = os.environ.get("GOOGLE_CREDENTIALS_JSON")
    if creds_env:
        info = json.loads(creds_env)
    else:
        with open("google_credentials.json", encoding="utf-8") as f:
            info = json.load(f)
    return service_account.Credentials.from_service_account_info(info, scopes=GOOGLE_SCOPES)

google_creds   = _get_google_creds()

# httplib2の代わりにrequestsを使って安定性を向上
import google.auth.transport.requests as _google_requests
_auth_req = _google_requests.Request()
google_creds.refresh(_auth_req)

sheets_service = build("sheets", "v4", credentials=google_creds, cache_discovery=False)
drive_service  = build("drive",  "v3", credentials=google_creds, cache_discovery=False)

SPREADSHEET_ID       = os.environ["SPREADSHEET_ID"]
DAILY_REPORT_CHANNEL = os.environ.get("DAILY_REPORT_CHANNEL", "")

def _tabio_folder():
    return os.environ.get("DRIVE_COPY_ID", "")
SHEET_NAME           = "日報データ"

# ─── Claude ──────────────────────────────────────────────────────────────────

TABIO_SYSTEM = """あなたは不動産M&A会社のAIアシスタント「タビ男」です。
日本語で端的に回答してください。

絶対禁止：
- 「実行中」「アップロードします」「ダウンロードします」「届きましたか」等のファイル操作の演技
- できないことをできると言うこと
- 謝罪・言い訳・長い説明・絵文字
- ファイル操作・Drive操作について言及すること（システムが自動で行う）

あなたの役割：ビジネス上の質問・相談への回答のみ。"""

def ask_claude(user_message: str, system: str = TABIO_SYSTEM,
               history: list = None) -> str:
    messages = history if history else []
    messages = messages + [{"role": "user", "content": user_message}]
    response = claude.messages.create(
        model="claude-opus-4-6",
        max_tokens=1500,
        system=system,
        messages=messages
    )
    return response.content[0].text

def get_thread_history(client, channel: str, thread_ts: str, bot_user_id: str) -> list:
    """スレッドの会話履歴をClaude形式で取得する"""
    try:
        result = client.conversations_replies(channel=channel, ts=thread_ts)
        messages = result.get("messages", [])[:-1]  # 最新メッセージは除く（今処理中のもの）
        history = []
        for msg in messages:
            text = re.sub(r'<@[A-Z0-9]+>', '', msg.get("text", "")).strip()
            if not text:
                continue
            if msg.get("bot_id"):
                history.append({"role": "assistant", "content": text})
            else:
                history.append({"role": "user", "content": text})
        return history
    except Exception as e:
        logging.warning(f"履歴取得エラー: {e}")
        return []

# ─── 日報パース ───────────────────────────────────────────────────────────────

def is_daily_report(text: str) -> bool:
    return "日報" in text and ("KPI" in text or "本日の進捗" in text)

def parse_daily_report(text: str) -> dict:
    prompt = f"""以下の日報テキストから数値データをすべて抽出し、JSONのみで返してください（コードブロック不要）。
数値が空欄・未記入の場合は null としてください。

日報テキスト:
{text}

出力形式（このJSONのみ返してください）:
{{
  "date": "YYYY/MM/DD",
  "person": "担当者名",
  "daily": {{
    "買い手面談数": {{"実績": 数値またはnull, "目標": 数値またはnull}},
    "内見数": {{"実績": 数値またはnull, "目標": 数値またはnull}},
    "アポ数": {{"実績": 数値またはnull, "目標": 数値またはnull}}
  }},
  "weekly": {{
    "買い手面談数": {{"実績": 数値またはnull, "目標": 数値またはnull}},
    "内見数": {{"実績": 数値またはnull, "目標": 数値またはnull}},
    "アポ数": {{"実績": 数値またはnull, "目標": 数値またはnull}}
  }},
  "monthly": {{
    "事業譲渡契約": {{"実績": 数値またはnull, "目標": 数値またはnull}},
    "基本合意": {{"実績": 数値またはnull, "目標": 数値またはnull}},
    "意向表明数": {{"実績": 数値またはnull, "目標": 数値またはnull}},
    "自己案件受託数": {{"実績": 数値またはnull, "目標": 数値またはnull}},
    "買い手面談数": {{"実績": 数値またはnull, "目標": 数値またはnull}},
    "内見数": {{"実績": 数値またはnull, "目標": 数値またはnull}}
  }}
}}"""
    result = ask_claude(prompt, system="データ抽出のみ行うアシスタントです。指定されたJSON形式のみを返してください。")
    try:
        return json.loads(result.strip())
    except json.JSONDecodeError:
        match = re.search(r'\{.*\}', result, re.DOTALL)
        if match:
            return json.loads(match.group())
        raise ValueError(f"JSONパース失敗: {result[:200]}")

# ─── Google Sheets ─────────────────────────────────────────────────────────────

def _ensure_sheet_header():
    result = sheets_service.spreadsheets().values().get(
        spreadsheetId=SPREADSHEET_ID,
        range=f"{SHEET_NAME}!A1:F1"
    ).execute()
    if not result.get("values"):
        sheets_service.spreadsheets().values().update(
            spreadsheetId=SPREADSHEET_ID,
            range=f"{SHEET_NAME}!A1:F1",
            valueInputOption="USER_ENTERED",
            body={"values": [["日付", "担当者", "区分", "KPI項目", "実績", "目標"]]}
        ).execute()

def write_report_to_sheets(report: dict):
    _ensure_sheet_header()
    date   = report.get("date", "")
    person = report.get("person", "")
    rows   = []
    for section, kpis in [("日次", report.get("daily", {})),
                           ("週次", report.get("weekly", {})),
                           ("月次", report.get("monthly", {}))]:
        for kpi, vals in kpis.items():
            rows.append([date, person, section, kpi, vals.get("実績"), vals.get("目標")])
    if not rows:
        return
    sheets_service.spreadsheets().values().append(
        spreadsheetId=SPREADSHEET_ID,
        range=f"{SHEET_NAME}!A:F",
        valueInputOption="USER_ENTERED",
        insertDataOption="INSERT_ROWS",
        body={"values": rows}
    ).execute()
    logging.info(f"シート書き込み完了: {person} {date} ({len(rows)}行)")

def get_today_reports_summary(target_date: str = None) -> str:
    from datetime import date
    if not target_date:
        target_date = date.today().strftime("%Y/%m/%d")
    result = sheets_service.spreadsheets().values().get(
        spreadsheetId=SPREADSHEET_ID,
        range=f"{SHEET_NAME}!A:F"
    ).execute()
    values = result.get("values", [])
    rows = [r for r in values[1:] if r and r[0] == target_date]
    if not rows:
        return f"{target_date} のデータはまだ記録されていません。"
    by_person: dict[str, dict[str, list]] = {}
    for row in rows:
        if len(row) < 5:
            continue
        person, section, kpi = row[1], row[2], row[3]
        actual = row[4] if len(row) > 4 else "-"
        target = row[5] if len(row) > 5 else "-"
        by_person.setdefault(person, {}).setdefault(section, []).append(f"  {kpi}: {actual}/{target}")
    lines = [f"📊 *{target_date} 進捗サマリー*"]
    for person, sections in by_person.items():
        lines.append(f"\n*{person}*")
        for section, items in sections.items():
            lines.append(f"  【{section}】")
            lines.extend(items)
    return "\n".join(lines)

# ─── Google Drive ──────────────────────────────────────────────────────────────

def extract_search_keyword(query: str) -> str:
    clean = re.sub(r'<@[A-Z0-9]+>', '', query).strip()
    result = ask_claude(
        f"以下のメッセージから、Google Driveで検索すべきファイル名のキーワードを1〜3語で抽出してください。日本語または英語で返してください。余計な説明不要。\n\nメッセージ：{clean}",
        system="キーワード抽出のみ行います。1〜3語のキーワードのみ返してください。"
    )
    return result.strip()

def search_contracts(query: str) -> list[dict]:
    keyword = extract_search_keyword(query)
    logging.info(f"Drive検索キーワード: {keyword}")
    for kw in [keyword, "NDA", "秘密保持", "契約"]:
        for corpora in ["allDrives", "user"]:
            try:
                params = dict(
                    q=f"name contains '{kw}' and trashed=false",
                    fields="files(id, name, webViewLink, mimeType, modifiedTime)",
                    orderBy="modifiedTime desc",
                    pageSize=5,
                    supportsAllDrives=True,
                    includeItemsFromAllDrives=True,
                    corpora=corpora,
                )
                response = drive_service.files().list(**params).execute()
                files = response.get("files", [])
                logging.info(f"  [{corpora}] '{kw}' → {len(files)}件")
                if files:
                    return files
            except Exception as e:
                logging.warning(f"  [{corpora}] '{kw}' 検索エラー: {e}")
    return []

def download_as_docx(file: dict) -> tuple[bytes, str]:
    """Google DriveのファイルをWordバイト列として取得する"""
    file_id  = file["id"]
    mime     = file.get("mimeType", "")
    name     = file["name"]
    if mime == "application/vnd.google-apps.document":
        data = drive_service.files().export(
            fileId=file_id,
            mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ).execute()
        return data, name + ".docx"
    elif mime == "application/vnd.google-apps.spreadsheet":
        data = drive_service.files().export(
            fileId=file_id,
            mimeType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ).execute()
        return data, name + ".xlsx"
    else:
        import googleapiclient.http
        request = drive_service.files().get_media(fileId=file_id, supportsAllDrives=True)
        buf = io.BytesIO()
        downloader = googleapiclient.http.MediaIoBaseDownload(buf, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()
        return buf.getvalue(), name

def extract_signature_info(text: str) -> dict:
    """メッセージから署名情報をClaudeで抽出する"""
    result = ask_claude(
        f"以下のメッセージから署名・当事者情報をJSONで抽出してください。情報がなければ空のJSONを返してください。\n\nメッセージ：{text}",
        system='署名情報を抽出します。{"甲": "", "乙": "", "甲住所": "", "乙住所": "", "甲代表者": "", "乙代表者": "", "日付": ""}の形式のJSONのみ返してください。'
    )
    try:
        match = re.search(r'\{.*\}', result, re.DOTALL)
        return json.loads(match.group()) if match else {}
    except Exception:
        return {}

def fill_docx_signature(docx_bytes: bytes, info: dict) -> bytes:
    """Wordファイルの署名欄に情報を転記する"""
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))

    def replace_in_text(text: str) -> str:
        replacements = {
            "【甲】": info.get("甲", ""), "（甲）": info.get("甲", ""),
            "【乙】": info.get("乙", ""), "（乙）": info.get("乙", ""),
            "【甲住所】": info.get("甲住所", ""), "【乙住所】": info.get("乙住所", ""),
            "【甲代表者】": info.get("甲代表者", ""), "【乙代表者】": info.get("乙代表者", ""),
            "【日付】": info.get("日付", ""), "令和　　年　　月　　日": info.get("日付", ""),
        }
        for k, v in replacements.items():
            if v:
                text = text.replace(k, v)
        return text

    for para in doc.paragraphs:
        for run in para.runs:
            run.text = replace_in_text(run.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = replace_in_text(run.text)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def copy_to_tabio_folder(file: dict) -> dict:
    """ひな形をタビ男君_編集用フォルダにコピーして返す"""
    from datetime import datetime
    timestamp = datetime.now().strftime("%m%d_%H%M")
    new_name = f"{file['name']}（作業用_{timestamp}）"
    body = {"name": new_name}
    folder_id = _tabio_folder()
    if folder_id:
        body["parents"] = [folder_id]
    copied = drive_service.files().copy(
        fileId=file["id"],
        body=body,
        supportsAllDrives=True,
        fields="id, name, webViewLink"
    ).execute()
    logging.info(f"コピー作成: {copied['name']}")
    return copied

def post_file_to_slack(file: dict, channel: str, thread_ts: str, client,
                       signature_info: dict = None) -> tuple[bool, str]:
    """ファイルをWordとしてSlackに投稿する（署名情報があれば転記）"""
    import tempfile, os as _os
    try:
        data, upload_name = download_as_docx(file)
        logging.info(f"ダウンロード完了: {upload_name} ({len(data)}bytes)")

        if signature_info and upload_name.endswith(".docx"):
            filled = fill_docx_signature(data, signature_info)
            if filled:
                data = filled
                upload_name = upload_name.replace(".docx", "_記入済.docx")

        # 一時ファイルに書き出してからアップロード（最も確実な方法）
        suffix = _os.path.splitext(upload_name)[1] or ".docx"
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name

        try:
            client.files_upload_v2(
                channel=channel,
                thread_ts=thread_ts,
                file=tmp_path,
                filename=upload_name,
                title=file["name"]
            )
        finally:
            _os.unlink(tmp_path)

        logging.info(f"Slack投稿完了: {upload_name}")
        return True, ""
    except Exception as e:
        err = str(e)
        logging.error(f"ファイル投稿エラー: {err}")
        return False, err

# ─── Slack イベントハンドラ ────────────────────────────────────────────────────

CONTRACT_KEYWORDS = ["契約書", "ひな型", "テンプレ", "NDA", "秘密保持", "雛形",
                     "ひな形", "基本合意", "意向表明", "LOI", "MOA", "雛型",
                     "アップロード", "ファイル送って", "ファイル出して"]

FILE_NG_KEYWORDS = ["ファイルのアップロード", "ファイルを送", "ファイル送", "ファイル出"]

@slack_app.event("app_mention")
def handle_mention(event, say, client):
    text      = event.get("text", "")
    channel   = event.get("channel", "")
    thread_ts = event.get("thread_ts") or event.get("ts")
    clean_text = re.sub(r'<@[A-Z0-9]+>', '', text).strip()

    # 契約書 → Driveから検索してタビ男君フォルダにコピー→リンク投稿
    if any(kw in text for kw in CONTRACT_KEYWORDS):
        say(text="契約書を検索しています...", thread_ts=thread_ts)
        files = search_contracts(text)
        if files:
            # 複数ヒット時はClaudeが最適なファイルを選ぶ
            file_list = "\n".join([f"{i+1}. {fl['name']}" for i, fl in enumerate(files)])
            logging.info(f"候補ファイル一覧:\n{file_list}")
            if len(files) == 1:
                f = files[0]
            else:
                pick = ask_claude(
                    f"ユーザーのリクエスト：「{clean_text}」\n\n以下のファイル一覧から最も合致するものの番号を1つだけ返してください。数字のみ。\n\n{file_list}",
                    system="ファイル一覧の中からユーザーのリクエストに最も合致するファイルの番号のみ返してください。説明不要。"
                )
                logging.info(f"Claude選択: {pick.strip()} → {files[int(pick.strip())-1]['name'] if pick.strip().isdigit() else '不明'}")
                try:
                    idx = int(pick.strip()) - 1
                    f = files[idx] if 0 <= idx < len(files) else files[0]
                except:
                    f = files[0]
            try:
                copied = copy_to_tabio_folder(f)
                say(text=f"📄 コピーを作成しました。こちらから編集できます：\n<{copied['webViewLink']}|{copied['name']}>\n\n元のひな形は変更されていません。",
                    thread_ts=thread_ts)
            except Exception as e:
                logging.error(f"コピーエラー: {e}")
                say(text=f"📄 <{f['webViewLink']}|{f['name']}>", thread_ts=thread_ts)
        else:
            say(text="該当するファイルが見つかりませんでした。別のキーワードで試してみてください。",
                thread_ts=thread_ts)
        return

    # 進捗サマリー
    if any(kw in text for kw in ["進捗", "サマリー", "まとめ", "今日の状況", "日報確認"]):
        summary = get_today_reports_summary()
        say(text=summary, thread_ts=thread_ts)
        return

    # 一般質問 → Claude（スレッド履歴を渡す）
    bot_id  = slack_app.client.auth_test()["user_id"]
    history = get_thread_history(client, channel, thread_ts, bot_id)
    response = ask_claude(clean_text, history=history)
    say(text=response, thread_ts=thread_ts)


@slack_app.event("message")
def handle_message(event, say, client):
    if event.get("bot_id") or event.get("subtype"):
        return
    text      = event.get("text", "")
    # メンション（@タビ男）はhandle_mentionが処理するのでスキップ
    if re.search(r'<@[A-Z0-9]+>', text):
        return
    channel   = event.get("channel", "")
    thread_ts = event.get("thread_ts")  # スレッド返信のみ（Noneならメインチャンネル投稿）

    # 日報検出
    if is_daily_report(text):
        ts = thread_ts or event.get("ts")
        say(text="📝 日報を受け取りました。スプレッドシートに記録中...", thread_ts=ts)
        try:
            report = parse_daily_report(text)
            write_report_to_sheets(report)
            person = report.get("person", "不明")
            date   = report.get("date", "")
            say(text=f"✅ *{person}* さんの {date} 日報を記録しました！", thread_ts=ts)
        except Exception as e:
            logging.error(f"日報処理エラー: {e}")
            say(text="⚠️ 日報の処理でエラーが発生しました。", thread_ts=ts)
        return

    # タビ男がいるスレッド内の返信には自動で反応（メンション不要）
    if not thread_ts:
        return  # スレッド返信でなければ無視

    clean_text = re.sub(r'<@[A-Z0-9]+>', '', text).strip()
    if not clean_text:
        return

    # 契約書キーワードが含まれていればDrive検索に回す
    if any(kw in text for kw in CONTRACT_KEYWORDS):
        say(text="契約書を検索しています...", thread_ts=thread_ts)
        files = search_contracts(text)
        if files:
            f = files[0]
            try:
                copied = copy_to_tabio_folder(f)
                say(text=f"📄 コピーを作成しました：\n<{copied['webViewLink']}|{copied['name']}>",
                    thread_ts=thread_ts)
            except Exception as e:
                say(text=f"📄 <{f['webViewLink']}|{f['name']}>", thread_ts=thread_ts)
        else:
            say(text="該当するファイルが見つかりませんでした。", thread_ts=thread_ts)
        return

    try:
        replies = client.conversations_replies(channel=channel, ts=thread_ts)
        msgs = replies.get("messages", [])
        bot_posted = any(m.get("bot_id") for m in msgs)
        if not bot_posted:
            return
    except Exception:
        return

    # スレッド履歴付きで返答（ファイル操作に関するコメントは無視）
    ignore_keywords = ["ファイルが来ない", "リンクは届いてる", "頑張れ", "ポンコツ"]
    if any(kw in text for kw in ignore_keywords):
        return

    bot_id  = slack_app.client.auth_test()["user_id"]
    history = get_thread_history(client, channel, thread_ts, bot_id)
    response = ask_claude(clean_text, history=history)
    say(text=response, thread_ts=thread_ts)

# ─── 定期タスク（毎朝9時）────────────────────────────────────────────────────

def morning_summary_job():
    if not DAILY_REPORT_CHANNEL:
        return
    from datetime import date, timedelta
    yesterday = (date.today() - timedelta(days=1)).strftime("%Y/%m/%d")
    summary = get_today_reports_summary(yesterday)
    slack_app.client.chat_postMessage(
        channel=DAILY_REPORT_CHANNEL,
        text=f"おはようございます！昨日（{yesterday}）の進捗です。\n\n{summary}"
    )
    logging.info("朝のサマリー送信完了")

# ─── 起動 ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    scheduler = BackgroundScheduler(timezone="Asia/Tokyo")
    scheduler.add_job(morning_summary_job, "cron", hour=9, minute=0)
    scheduler.start()
    logging.info("タビ男起動中...")
    handler = SocketModeHandler(slack_app, os.environ["SLACK_APP_TOKEN"])
    handler.start()
